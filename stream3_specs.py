# ─────────────────────────────────────────────────────────────
# stream3_specs.py — Technical Specifications (.docx / .pdf)
# Usage: python stream3_specs.py
# ─────────────────────────────────────────────────────────────

import json
import logging
import os
from pathlib import Path
from typing import List, Dict, Tuple
from tqdm import tqdm

import docx
import pdfplumber
import fitz

from config import TECH_SPEC_DIR, SPEC_JSONL, SPEC_REJECT_CSV
from utils import make_content_hash, clean_text_field
from stream2_manuals import (
    is_digital_pdf,
    extract_text_from_digital_pdf,
    extract_text_from_scanned_pdf,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)s: %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("logs/stream3.log", encoding="utf-8")
    ]
)
log = logging.getLogger(__name__)


def extract_text_tables_and_metadata_from_docx(
    docx_path: str
) -> Tuple[List[str], List[List[List[str]]], Dict]:
    """Extract paragraphs, tables, and metadata from a .docx file."""
    try:
        doc = docx.Document(docx_path)
        text_paragraphs = []
        tables = []
        metadata = {
            "title"       : doc.core_properties.title,
            "author"      : doc.core_properties.author,
            "created"     : str(doc.core_properties.created),
            "modified"    : str(doc.core_properties.modified),
            "revision"    : doc.core_properties.revision,
        }

        for para in doc.paragraphs:
            cleaned = clean_text_field(para.text)
            if cleaned:
                text_paragraphs.append(cleaned)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [clean_text_field(cell.text) for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return text_paragraphs, tables, metadata
    except Exception as e:
        log.error(f"Error processing DOCX {docx_path}: {e}")
        return [], [], {}


def extract_tables_from_pdf(pdf_path: str) -> List[List[List[str]]]:
    """Extract all tables from a PDF using pdfplumber."""
    all_tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    cleaned_table = [
                        [clean_text_field(cell) if cell is not None else "" for cell in row]
                        for row in table
                    ]
                    if cleaned_table:
                        all_tables.append(cleaned_table)
    except Exception as e:
        log.error(f"Error extracting tables from PDF {pdf_path}: {e}")
    return all_tables


def chunk_pdf_text_into_sections(
    text_pages: List[str], source_filename: str
) -> List[Dict]:
    """Chunk PDF text by paragraph, skip very short ones."""
    sections = []
    full_text = "\n\n".join(text_pages)
    paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
    chunk_id = 0
    for para in paragraphs:
        if len(para.split()) < 10:
            continue
        content = clean_text_field(para)
        if content:
            sections.append({
                "id"          : f"SPEC-PDF-{source_filename.replace('.', '_')}-text-{chunk_id:05d}",
                "doc_type"    : "technical_spec_text_chunk",
                "source_file" : source_filename,
                "content"     : content,
                "content_hash": make_content_hash(content),
                "metadata"    : {},
            })
            chunk_id += 1
    return sections


def run_stream3(spec_dir: Path, output_jsonl: Path, reject_csv: Path):
    log.info("=" * 60)
    log.info("STREAM 3 — Technical Specifications (.docx / .pdf)")
    log.info("=" * 60)

    spec_files = [
        f for f in spec_dir.iterdir()
        if f.suffix.lower() in (".docx", ".pdf")
    ]

    if not spec_files:
        log.warning(f"No .docx or .pdf files found in {spec_dir}")
        return

    log.info(f"Found {len(spec_files)} specification file(s)")
    accepted = []
    rejected = []
    seen_hashes = set()

    for file_path in tqdm(spec_files, desc="Processing spec files"):
        file_name = file_path.name
        ext = file_path.suffix.lower()
        file_meta = {"source_file_path": str(file_path)}
        processed_count = 0

        if ext == ".docx":
            paragraphs, tables, docx_meta = extract_text_tables_and_metadata_from_docx(str(file_path))
            file_meta.update(docx_meta)

            for i, text in enumerate(paragraphs):
                if not text or len(text) < 30:
                    continue
                record = {
                    "id"         : f"SPEC-DOCX-{file_name.replace('.', '_')}-text-{i:05d}",
                    "doc_type"   : "technical_spec_text",
                    "source_file": file_name,
                    "content"    : text,
                    "metadata"   : file_meta.copy(),
                }
                h = make_content_hash(record["content"])
                if h in seen_hashes:
                    record["reject_reason"] = "Duplicate content"
                    rejected.append(record)
                else:
                    seen_hashes.add(h)
                    record["content_hash"] = h
                    accepted.append(record)
                    processed_count += 1

            for i, table_data in enumerate(tables):
                if not table_data:
                    continue
                table_str = "\n".join(["|" + "|".join(row) + "|" for row in table_data])
                record = {
                    "id"             : f"SPEC-DOCX-{file_name.replace('.', '_')}-table-{i:05d}",
                    "doc_type"       : "technical_spec_table",
                    "source_file"    : file_name,
                    "content"        : table_str,
                    "structured_data": table_data,
                    "metadata"       : file_meta.copy(),
                }
                h = make_content_hash(record["content"])
                if h in seen_hashes:
                    record["reject_reason"] = "Duplicate content"
                    rejected.append(record)
                else:
                    seen_hashes.add(h)
                    record["content_hash"] = h
                    accepted.append(record)
                    processed_count += 1

        elif ext == ".pdf":
            # Get PDF metadata
            try:
                doc = fitz.open(str(file_path))
                file_meta.update({
                    "title"      : doc.metadata.get("title"),
                    "author"     : doc.metadata.get("author"),
                    "page_count" : doc.page_count,
                })
                doc.close()
            except Exception as e:
                log.warning(f"Could not extract PDF metadata from {file_name}: {e}")

            if is_digital_pdf(str(file_path)):
                pages = extract_text_from_digital_pdf(str(file_path))
            else:
                pages = extract_text_from_scanned_pdf(str(file_path))

            pdf_tables = extract_tables_from_pdf(str(file_path))

            if not pages and not pdf_tables:
                log.warning(f"No content from {file_name} — skipping")
                rejected.append({"source_file": file_name,
                                  "reject_reason": "No text or tables extracted",
                                  "metadata": file_meta})
                continue

            for record in chunk_pdf_text_into_sections(pages, file_name):
                record["metadata"].update(file_meta)
                h = make_content_hash(record["content"])
                if h in seen_hashes:
                    record["reject_reason"] = "Duplicate content"
                    rejected.append(record)
                else:
                    seen_hashes.add(h)
                    record["content_hash"] = h
                    accepted.append(record)
                    processed_count += 1

            for i, table_data in enumerate(pdf_tables):
                table_str = "\n".join(["|" + "|".join(row) + "|" for row in table_data])
                record = {
                    "id"             : f"SPEC-PDF-{file_name.replace('.', '_')}-table-{i:05d}",
                    "doc_type"       : "technical_spec_table",
                    "source_file"    : file_name,
                    "content"        : table_str,
                    "structured_data": table_data,
                    "metadata"       : file_meta.copy(),
                }
                h = make_content_hash(record["content"])
                if h in seen_hashes:
                    record["reject_reason"] = "Duplicate content"
                    rejected.append(record)
                else:
                    seen_hashes.add(h)
                    record["content_hash"] = h
                    accepted.append(record)
                    processed_count += 1

        if processed_count == 0:
            if not any(r.get("source_file") == file_name for r in rejected):
                rejected.append({"source_file": file_name,
                                  "reject_reason": "No valid content extracted",
                                  "metadata": file_meta})

    # ── Write outputs ──────────────────────────────────────────
    with open(output_jsonl, "w", encoding="utf-8") as f:
        for rec in accepted:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")

    if rejected:
        import pandas as pd
        pd.DataFrame([{
            "source_file"  : r.get("source_file"),
            "reject_reason": r.get("reject_reason"),
            "metadata"     : json.dumps(r.get("metadata", {})),
        } for r in rejected]).to_csv(reject_csv, index=False)

    log.info("")
    log.info("─" * 40)
    log.info(f"✓ Accepted: {len(accepted)} chunks → {output_jsonl}")
    log.info(f"✗ Rejected: {len(rejected)} → {reject_csv}")
    log.info("─" * 40)


# ── Entry point ───────────────────────────────────────────────
if __name__ == "__main__":
    run_stream3(TECH_SPEC_DIR, SPEC_JSONL, SPEC_REJECT_CSV)
    log.info("\nStream 3 complete.")
